
from __future__ import annotations

import io
import os
from dataclasses import dataclass

import docx
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from googleapiclient.http import MediaIoBaseDownload
from pypdf import PdfReader
from tenacity import retry, retry_if_exception_type, stop_after_attempt, wait_exponential

from agentic_rag.exceptions import DriveAccessError
from agentic_rag.logging_config import get_logger
from agentic_rag.settings import Settings, get_settings

logger = get_logger(__name__)

SCOPES = ["https://www.googleapis.com/auth/drive.readonly"]

GOOGLE_EXPORT_MIME = {
    "application/vnd.google-apps.document": ("text/plain", ".txt"),
    "application/vnd.google-apps.spreadsheet": ("text/csv", ".csv"),
    "application/vnd.google-apps.presentation": ("text/plain", ".txt"),
}

SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".txt", ".md")


@dataclass(frozen=True)
class DriveDocument:
    id: str
    name: str
    text: str
    source: str


class DriveReader:
    def __init__(self, settings: Settings | None = None):
        self.settings = settings or get_settings()
        self._service = None

    @property
    def service(self):
        if self._service is None:
            sa_file = self.settings.GOOGLE_SERVICE_ACCOUNT_FILE
            if not os.path.exists(sa_file):
                raise DriveAccessError(
                    f"Service account file not found at '{sa_file}'. "
                    f"See README.md for setup instructions."
                )
            creds = service_account.Credentials.from_service_account_file(sa_file, scopes=SCOPES)
            self._service = build("drive", "v3", credentials=creds, cache_discovery=False)
        return self._service

    @retry(
        retry=retry_if_exception_type(HttpError),
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=1, max=8),
        reraise=True,
    )
    def list_page(self, folder_id: str, page_token: str | None):
        return (
            self.service.files()
            .list(
                q=f"'{folder_id}' in parents and trashed = false",
                fields="nextPageToken, files(id, name, mimeType)",
                pageToken=page_token,
            )
            .execute()
        )

    def list_files(self, folder_id: str | None = None) -> list[dict]:
        folder_id = folder_id or self.settings.GOOGLE_DRIVE_FOLDER_ID
        if not folder_id:
            raise DriveAccessError("No Google Drive folder ID configured (GOOGLE_DRIVE_FOLDER_ID).")

        files: list[dict] = []
        page_token = None
        try:
            while True:
                resp = self.list_page(folder_id, page_token)
                for f in resp.get("files", []):
                    if f["mimeType"] == "application/vnd.google-apps.folder":
                        files.extend(self.list_files(f["id"]))
                    else:
                        files.append(f)
                page_token = resp.get("nextPageToken")
                if not page_token:
                    break
        except HttpError as e:
            raise DriveAccessError(f"Failed to list Drive folder '{folder_id}': {e}") from e
        return files

    @retry(
        retry=retry_if_exception_type(HttpError),
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=1, max=8),
        reraise=True,
    )
    def download_bytes(self, file_id: str, mime_type: str) -> bytes:
        if mime_type in GOOGLE_EXPORT_MIME:
            export_mime, _ = GOOGLE_EXPORT_MIME[mime_type]
            request = self.service.files().export_media(fileId=file_id, mimeType=export_mime)
        else:
            request = self.service.files().get_media(fileId=file_id)

        buf = io.BytesIO()
        downloader = MediaIoBaseDownload(buf, request)
        done = False
        while not done:
            _, done = downloader.next_chunk()
        return buf.getvalue()

    @staticmethod
    def extract_text(raw: bytes, name: str, mime_type: str) -> str:
        try:
            if mime_type in GOOGLE_EXPORT_MIME or mime_type.startswith("text/"):
                return raw.decode("utf-8", errors="ignore")
            if mime_type == "application/pdf" or name.lower().endswith(".pdf"):
                reader = PdfReader(io.BytesIO(raw))
                return "\n".join(page.extract_text() or "" for page in reader.pages)
            if name.lower().endswith(".docx"):
                d = docx.Document(io.BytesIO(raw))
                return "\n".join(p.text for p in d.paragraphs)
        except Exception as e:  # extraction failures shouldn't abort the whole ingestion
            logger.warning("Could not extract text from '%s': %s", name, e)
            return ""
        return ""

    def is_supported(self, f: dict) -> bool:
        mime_type = f["mimeType"]
        return (
            mime_type in GOOGLE_EXPORT_MIME
            or mime_type.startswith("text/")
            or mime_type == "application/pdf"
            or f["name"].lower().endswith(SUPPORTED_EXTENSIONS)
        )

    def fetch_documents(self) -> list[DriveDocument]:
        docs: list[DriveDocument] = []
        files = self.list_files()
        logger.info("Found %d files in Drive folder; filtering supported types...", len(files))

        for f in files:
            if not self.is_supported(f):
                continue
            try:
                raw = self.download_bytes(f["id"], f["mimeType"])
                text = self.extract_text(raw, f["name"], f["mimeType"])
                if text.strip():
                    docs.append(
                        DriveDocument(
                            id=f["id"], name=f["name"], text=text, source=f"gdrive://{f['name']}"
                        )
                    )
            except HttpError as e:
                logger.warning("Skipped '%s' due to API error: %s", f["name"], e)
            except Exception as e:
                logger.warning("Skipped '%s' due to unexpected error: %s", f["name"], e)

        logger.info("Extracted text from %d/%d Drive files.", len(docs), len(files))
        return docs
